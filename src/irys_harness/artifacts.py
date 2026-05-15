from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook

ILLEGAL_XML_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
XLSX_CELL_CHAR_LIMIT = 500


def render_deliverables(
    *,
    output_dir: str | Path,
    deliverables: list[str],
    title: str,
    packet: dict[str, Any],
) -> list[dict[str, Any]]:
    root = Path(output_dir)
    root.mkdir(parents=True, exist_ok=True)
    apply_deliverable_coverage_audit(packet, deliverables)
    artifacts: list[dict[str, Any]] = []
    for filename in deliverables:
        path = root / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        if path.suffix.lower() == ".xlsx":
            render_xlsx(path, title=title, deliverable=filename, packet=packet)
        else:
            render_docx(path, title=title, deliverable=filename, packet=packet, deliverables=deliverables)
        artifacts.append(
            {
                "filename": filename,
                "path": str(path),
                "exists": path.exists(),
                "bytes": path.stat().st_size if path.exists() else 0,
            }
        )
    return artifacts


def render_docx(
    path: Path,
    *,
    title: str,
    deliverable: str,
    packet: dict[str, Any],
    deliverables: list[str] | None = None,
) -> None:
    doc = Document()
    doc.add_heading(safe_text(title), level=1)
    doc.add_paragraph(safe_text(f"Deliverable: {deliverable}"))
    deliverable_plan = find_deliverable_plan(packet, deliverable)
    if deliverable_plan:
        doc.add_heading("Deliverable Contract", level=2)
        for section in deliverable_plan.get("required_sections", [])[:12]:
            doc.add_paragraph(safe_text(section), style=None)
    draft_answer = packet.get("draft_answer")
    if draft_answer:
        doc.add_heading("Draft Work Product", level=2)
        if is_probably_encoded_artifact(str(draft_answer)):
            draft_answer = packet.get("plain_text_fallback") or packet.get("cheap_worker_summary") or ""
        draft_answer = extract_deliverable_draft_text(
            str(draft_answer),
            deliverable=deliverable,
            deliverables=deliverables or [deliverable],
        )
        for block in str(draft_answer).split("\n\n"):
            if block.strip():
                doc.add_paragraph(safe_text(block.strip()))
    coverage = find_coverage_record(packet, deliverable)
    if coverage and not coverage.get("present_in_draft") and coverage.get("fallback_text"):
        doc.add_heading("Deliverable-Specific Coverage Fill", level=2)
        for block in str(coverage["fallback_text"]).split("\n\n"):
            if block.strip():
                doc.add_paragraph(safe_text(block.strip()))
    appendix = packet.get("artifact_appendix")
    if appendix:
        doc.add_heading("Structured Findings Appendix", level=2)
        for block in str(appendix).split("\n\n"):
            if block.strip():
                doc.add_paragraph(safe_text(block.strip()))
    doc.add_heading("Candidate Evidence Packet", level=2)
    for item in packet.get("verified_evidence", []):
        claim = item.get("claim", "")
        support = item.get("raw_support", "")
        source = item.get("source", {})
        doc.add_paragraph(safe_text(f"Claim: {claim}"))
        doc.add_paragraph(safe_text(f"Source: {source.get('doc_id')} / {source.get('chunk_id')}"))
        doc.add_paragraph(safe_text(support, limit=1200))
    unresolved = packet.get("unresolved", [])
    if unresolved:
        doc.add_heading("Unresolved", level=2)
        for item in unresolved:
            doc.add_paragraph(safe_text(item))
    doc.save(path)


def render_xlsx(path: Path, *, title: str, deliverable: str, packet: dict[str, Any]) -> None:
    workbook = Workbook()
    deliverable_plan = find_deliverable_plan(packet, deliverable)
    sheet_plans = normalize_sheet_plans(deliverable_plan)
    if "Evidence" not in {plan["name"] for plan in sheet_plans}:
        sheet_plans.append({"name": "Evidence", "purpose": "Source evidence and trace support"})

    draft_text = str(packet.get("draft_answer") or "")
    if draft_text:
        draft_text = extract_deliverable_draft_text(
            draft_text,
            deliverable=deliverable,
            deliverables=packet_deliverable_names(packet, fallback=[deliverable]),
        )
    text = "\n\n".join(
        item
        for item in [draft_text, str(packet.get("cheap_worker_summary") or "")]
        if item
    )
    tables = extract_markdown_tables(text)

    for index, sheet_plan in enumerate(sheet_plans):
        sheet = workbook.active if index == 0 else workbook.create_sheet()
        sheet.title = safe_sheet_title(sheet_plan["name"], used=workbook.sheetnames)
        render_planned_sheet(
            sheet,
            title=title,
            deliverable=deliverable,
            sheet_plan=sheet_plan,
            packet=packet,
            tables=tables,
            source_text=text,
        )
    workbook.save(path)


def extract_deliverable_draft_text(text: str, *, deliverable: str, deliverables: list[str]) -> str:
    """Return the filename-labeled section for a deliverable when synthesis emitted one."""
    if not text.strip():
        return text
    lines = text.splitlines()
    headings: list[tuple[int, str]] = []
    for index, line in enumerate(lines):
        matched = match_deliverable_heading(line, deliverables)
        if matched:
            headings.append((index, matched))
    start_indexes = [index for index, matched in headings if matched.lower() == deliverable.lower()]
    if not start_indexes:
        return text
    start = start_indexes[0]
    end = len(lines)
    for index, matched in headings:
        if index > start and matched.lower() != deliverable.lower():
            end = index
            break
    section = "\n".join(lines[start:end]).strip()
    return section or text


def match_deliverable_heading(line: str, deliverables: list[str]) -> str | None:
    normalized = normalize_deliverable_heading(line)
    if not normalized or len(normalized) > 200:
        return None
    for deliverable in deliverables:
        name = deliverable.lower()
        patterns = [
            name,
            f"deliverable: {name}",
            f"file: {name}",
            f"filename: {name}",
            f"artifact: {name}",
        ]
        if normalized in patterns or any(normalized.startswith(f"{pattern} ") for pattern in patterns):
            return deliverable
    return None


def normalize_deliverable_heading(line: str) -> str:
    cleaned = line.strip()
    if not cleaned:
        return ""
    if len(cleaned) > 240:
        return ""
    cleaned = re.sub(r"^\s{0,3}(?:#{1,6}|[-*+>]|\d+[.)])\s*", "", cleaned)
    cleaned = cleaned.replace("**", "").replace("__", "").strip()
    cleaned = cleaned.strip("`'\" ")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip(":- ").lower()


def find_deliverable_plan(packet: dict[str, Any], deliverable: str) -> dict[str, Any] | None:
    contract = packet.get("deliverable_contract") or {}
    for item in contract.get("deliverables", []) or []:
        if str(item.get("filename", "")).lower() == deliverable.lower():
            return item
    return None


def packet_deliverable_names(packet: dict[str, Any], *, fallback: list[str] | None = None) -> list[str]:
    names: list[str] = []
    for item in packet.get("deliverables", []) or []:
        if isinstance(item, str):
            names.append(item)
    contract = packet.get("deliverable_contract") or {}
    for item in contract.get("deliverables", []) or []:
        filename = str(item.get("filename") or "")
        if filename:
            names.append(filename)
    if not names and fallback:
        names.extend(fallback)
    return dedupe_preserve_order(names)


def apply_deliverable_coverage_audit(packet: dict[str, Any], deliverables: list[str]) -> list[dict[str, Any]]:
    existing = packet.get("deliverable_coverage_audit")
    if isinstance(existing, list) and existing:
        return existing
    draft_answer = str(packet.get("draft_answer") or "")
    records: list[dict[str, Any]] = []
    for deliverable in deliverables:
        present, matched_by = deliverable_present_in_text(draft_answer, deliverable)
        plan = find_deliverable_plan(packet, deliverable)
        record: dict[str, Any] = {
            "filename": deliverable,
            "present_in_draft": present,
            "matched_by": matched_by,
            "artifact_role": plan.get("artifact_role") if plan else None,
            "required_sections": list(plan.get("required_sections", []) or []) if plan else [],
        }
        if not present:
            record["fallback_text"] = build_missing_deliverable_fill(deliverable, plan, packet)
        records.append(record)
    packet["deliverable_coverage_audit"] = records
    missing = [record["filename"] for record in records if not record["present_in_draft"]]
    if missing:
        packet["deliverable_coverage_summary"] = {
            "missing_deliverables": missing,
            "present_deliverables": [record["filename"] for record in records if record["present_in_draft"]],
        }
    return records


def deliverable_present_in_text(text: str, deliverable: str) -> tuple[bool, str | None]:
    lower_text = text.lower()
    lower_name = deliverable.lower()
    stem = Path(lower_name).stem
    variants = [
        lower_name,
        stem,
        stem.replace("-", " "),
        stem.replace("_", " "),
    ]
    for variant in dedupe_preserve_order([item for item in variants if item]):
        if variant in lower_text:
            return True, variant
    return False, None


def find_coverage_record(packet: dict[str, Any], deliverable: str) -> dict[str, Any] | None:
    for record in packet.get("deliverable_coverage_audit", []) or []:
        if str(record.get("filename", "")).lower() == deliverable.lower():
            return record
    return None


def build_missing_deliverable_fill(
    deliverable: str,
    plan: dict[str, Any] | None,
    packet: dict[str, Any],
) -> str:
    artifact_goal = str((plan or {}).get("artifact_goal") or f"Produce {deliverable}.")
    required_sections = [str(section) for section in (plan or {}).get("required_sections", []) or []]
    lines = [
        deliverable,
        "",
        "Coverage audit note",
        "The synthesis draft did not include a filename-level section for this requested deliverable. The renderer preserved the full global draft and adds this deliverable-specific fill so the requested artifact remains inspectable.",
        "",
        "Artifact goal",
        artifact_goal,
    ]
    if required_sections:
        lines.extend(["", "Required sections"])
        for section in required_sections[:16]:
            lines.append(f"- {section}")
    support = select_deliverable_support(deliverable, plan, packet)
    if support:
        lines.extend(["", "Source-state notes"])
        for item in support:
            lines.append(f"- {item}")
    else:
        lines.extend(
            [
                "",
                "Source-state notes",
                "- No deterministic source line matched this deliverable name. Review the global draft, structured findings appendix, and candidate evidence packet for source support.",
            ]
        )
    return "\n".join(lines)


def select_deliverable_support(
    deliverable: str,
    plan: dict[str, Any] | None,
    packet: dict[str, Any],
    *,
    max_items: int = 8,
) -> list[str]:
    source_text = "\n".join(
        str(packet.get(key) or "")
        for key in ["cheap_worker_summary", "artifact_appendix", "draft_answer"]
    )
    keywords = deliverable_keywords(deliverable, plan)
    selected: list[str] = []
    for raw_line in source_text.splitlines():
        line = " ".join(raw_line.strip().split())
        if len(line) < 24:
            continue
        lower_line = line.lower()
        if any(keyword in lower_line for keyword in keywords):
            selected.append(safe_text(line, limit=500))
        if len(selected) >= max_items:
            break
    return dedupe_preserve_order(selected)


def deliverable_keywords(deliverable: str, plan: dict[str, Any] | None) -> list[str]:
    raw_terms = [
        Path(deliverable.lower()).stem.replace("-", " "),
        str((plan or {}).get("artifact_role") or "").replace("_", " "),
        str((plan or {}).get("artifact_goal") or ""),
        " ".join(str(section) for section in (plan or {}).get("required_sections", []) or []),
    ]
    stopwords = {
        "and",
        "the",
        "for",
        "with",
        "from",
        "this",
        "that",
        "source",
        "citations",
        "document",
        "documents",
        "section",
        "sections",
        "deliverable",
    }
    terms: list[str] = []
    for text in raw_terms:
        for token in re.findall(r"[a-z][a-z0-9]{3,}", text.lower()):
            if token not in stopwords:
                terms.append(token)
    return dedupe_preserve_order(terms)


def dedupe_preserve_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    deduped: list[str] = []
    for item in items:
        if item in seen:
            continue
        seen.add(item)
        deduped.append(item)
    return deduped


def normalize_sheet_plans(deliverable_plan: dict[str, Any] | None) -> list[dict[str, str]]:
    raw_plans = []
    if deliverable_plan:
        raw_plans = deliverable_plan.get("workbook_sheets", []) or []
    if not raw_plans:
        raw_plans = [
            {"name": "Issue Matrix", "purpose": "Task-specific issues, calculations, or comparisons"},
            {"name": "Source Evidence", "purpose": "Evidence supporting the workbook"},
        ]
    plans: list[dict[str, str]] = []
    for item in raw_plans:
        if isinstance(item, str):
            plans.append({"name": item, "purpose": ""})
        else:
            plans.append(
                {
                    "name": str(item.get("name") or "Sheet"),
                    "purpose": str(item.get("purpose") or ""),
                }
            )
    return plans


def render_planned_sheet(
    sheet: Any,
    *,
    title: str,
    deliverable: str,
    sheet_plan: dict[str, str],
    packet: dict[str, Any],
    tables: list[dict[str, Any]],
    source_text: str,
) -> None:
    sheet.append(["Task", safe_cell(title)])
    sheet.append(["Deliverable", safe_cell(deliverable)])
    sheet.append(["Sheet Purpose", safe_cell(sheet_plan.get("purpose", ""))])
    sheet.append([])
    if sheet_plan["name"].lower() == "evidence":
        append_evidence_rows(sheet, packet)
        return

    matched_tables = match_tables_to_sheet(
        tables,
        sheet_plan["name"],
        sheet_plan.get("purpose", ""),
    )
    if matched_tables:
        for table in matched_tables[:3]:
            append_markdown_table(sheet, table)
            sheet.append([])
    else:
        append_issue_template(sheet, sheet_plan["name"])
    append_relevant_text(sheet, source_text, sheet_plan["name"])
    append_evidence_rows(sheet, packet, max_rows=10)


def append_issue_template(sheet: Any, sheet_name: str) -> None:
    lower = sheet_name.lower()
    if any(term in lower for term in ["calculation", "model", "limitation", "shift", "utilization"]):
        sheet.append(["Item", "Source Input", "Formula / Standard", "Result", "Conclusion", "Source"])
    elif any(term in lower for term in ["register", "inventory", "matrix", "comparison", "log"]):
        sheet.append(["Item", "Document / Party", "Requirement", "Finding", "Risk", "Source"])
    else:
        sheet.append(["Issue", "Finding", "Impact", "Remediation", "Source"])


def append_evidence_rows(sheet: Any, packet: dict[str, Any], *, max_rows: int | None = None) -> None:
    sheet.append(["Claim", "Doc ID", "Chunk ID", "Support"])
    evidence = packet.get("verified_evidence", [])
    if max_rows is not None:
        evidence = evidence[:max_rows]
    for item in evidence:
        source = item.get("source", {})
        sheet.append(
            [
                safe_cell(item.get("claim", "")),
                safe_cell(source.get("doc_id", "")),
                safe_cell(source.get("chunk_id", "")),
                safe_cell(item.get("raw_support", "")),
            ]
        )


def extract_markdown_tables(text: str) -> list[dict[str, Any]]:
    lines = text.splitlines()
    tables: list[dict[str, Any]] = []
    index = 0
    heading = ""
    while index < len(lines) - 1:
        heading_match = re.match(r"^\s*#{1,6}\s+(.+?)\s*$", lines[index])
        if heading_match:
            heading = heading_match.group(1).strip()
            index += 1
            continue
        if "|" not in lines[index] or not re.search(r"\|\s*:?-{2,}:?\s*(\||$)", lines[index + 1]):
            index += 1
            continue
        table_heading = infer_plain_table_heading(lines, index) or heading
        header = split_markdown_row(lines[index])
        rows: list[list[str]] = []
        index += 2
        while index < len(lines) and "|" in lines[index]:
            row = split_markdown_row(lines[index])
            if row:
                rows.append(row)
            index += 1
        if header and rows:
            tables.append({"heading": table_heading, "headers": header, "rows": rows})
    return tables


def infer_plain_table_heading(lines: list[str], table_index: int) -> str:
    for offset in range(table_index - 1, max(-1, table_index - 5), -1):
        candidate = lines[offset].strip()
        if not candidate:
            continue
        if "|" in candidate:
            return ""
        normalized = normalize_plain_heading(candidate)
        if normalized:
            return normalized
        return ""
    return ""


def normalize_plain_heading(line: str) -> str:
    cleaned = line.strip()
    cleaned = re.sub(r"^\s{0,3}(?:#{1,6}|[-*+>]|\d+[.)])\s*", "", cleaned)
    cleaned = cleaned.replace("**", "").replace("__", "").strip()
    cleaned = cleaned.strip("`'\" ")
    cleaned = re.sub(r"\s+", " ", cleaned)
    if not cleaned or len(cleaned) > 100:
        return ""
    if cleaned.endswith(".") or re.search(r"[,:;]$", cleaned):
        return ""
    if len(cleaned.split()) > 8:
        return ""
    return cleaned


def split_markdown_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [safe_cell(cell.strip()) for cell in stripped.split("|")]


def match_tables_to_sheet(
    tables: list[dict[str, Any]],
    sheet_name: str,
    sheet_purpose: str = "",
) -> list[dict[str, Any]]:
    sheet_tokens = keyword_tokens(f"{sheet_name} {sheet_purpose}")
    scored: list[tuple[int, dict[str, Any]]] = []
    for table in tables:
        heading = str(table.get("heading", ""))
        table_text = heading + " " + " ".join(table.get("headers", []))
        table_text += " " + " ".join(" ".join(row) for row in table.get("rows", [])[:4])
        table_tokens = keyword_tokens(table_text)
        score = fuzzy_token_overlap(sheet_tokens, table_tokens)
        heading_tokens = keyword_tokens(heading)
        score += 2 * fuzzy_token_overlap(sheet_tokens, heading_tokens)
        if normalize_match_text(sheet_name) and normalize_match_text(sheet_name) in normalize_match_text(heading):
            score += 80
        if score and any(
            term in table_text.lower()
            for term in ["deficiency", "candidate", "recommended action", "source", "formula", "calculation"]
        ):
            score += 2
        if score:
            scored.append((score, table))
    if scored:
        return [table for _, table in sorted(scored, key=lambda item: item[0], reverse=True)]
    return tables[:1]


def fuzzy_token_overlap(left: set[str], right: set[str]) -> int:
    score = 0
    for token in left:
        if token in right:
            score += 1
            continue
        if len(token) < 5:
            continue
        if any(other.startswith(token) or token.startswith(other) for other in right if len(other) >= 5):
            score += 1
    return score


def normalize_match_text(text: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9]+", " ", text.lower())).strip()


def append_markdown_table(sheet: Any, table: dict[str, Any]) -> None:
    headers = [safe_cell(item) for item in table.get("headers", [])]
    if headers:
        sheet.append(headers)
    for row in table.get("rows", [])[:80]:
        sheet.append([safe_cell(item) for item in row])


def append_relevant_text(sheet: Any, text: str, sheet_name: str) -> None:
    snippets = relevant_snippets(text, sheet_name)
    if not snippets:
        return
    sheet.append([])
    sheet.append(["Relevant Worker Text"])
    for snippet in snippets[:8]:
        sheet.append([safe_cell(snippet)])


def relevant_snippets(text: str, sheet_name: str) -> list[str]:
    tokens = keyword_tokens(sheet_name)
    sections = re.split(r"(?m)^#{1,4}\s+", text)
    matches: list[str] = []
    for section in sections:
        section = section.strip()
        if not section:
            continue
        if tokens.intersection(keyword_tokens(section[:800])):
            matches.append(section[:5000])
    return matches


def is_probably_encoded_artifact(text: str) -> bool:
    stripped = str(text or "").strip()
    lower = stripped[:4000].lower()
    if any(marker in lower for marker in ["<base64_file", "<content>", "<?xml", "pk\x03\x04"]):
        return True
    if re.match(r"^```(?:xml|base64|html)?\s*<", stripped, flags=re.IGNORECASE | re.DOTALL):
        return True
    base64ish = re.sub(r"\s+", "", stripped[:12000])
    return len(base64ish) > 4000 and bool(re.fullmatch(r"[A-Za-z0-9+/=]+", base64ish))


def keyword_tokens(text: str) -> set[str]:
    stop = {"and", "the", "for", "with", "from", "source", "evidence", "sheet"}
    return {
        token
        for token in re.findall(r"[a-z0-9]{3,}", text.lower())
        if token not in stop
    }


def safe_sheet_title(title: str, *, used: list[str]) -> str:
    cleaned = re.sub(r"[\[\]:*?/\\]", " ", title).strip() or "Sheet"
    cleaned = re.sub(r"\s+", " ", cleaned)[:31].strip() or "Sheet"
    candidate = cleaned
    suffix = 2
    while candidate in used:
        suffix_text = f" {suffix}"
        candidate = f"{cleaned[: 31 - len(suffix_text)]}{suffix_text}"
        suffix += 1
    return candidate


def safe_text(value: Any, *, limit: int | None = None) -> str:
    text = "" if value is None else str(value)
    text = ILLEGAL_XML_RE.sub("", text)
    if limit is not None:
        return text[:limit]
    return text


def safe_cell(value: Any, *, limit: int = XLSX_CELL_CHAR_LIMIT) -> str:
    return safe_text(value, limit=limit)
