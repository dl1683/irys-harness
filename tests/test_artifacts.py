from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from irys_harness.artifacts import (
    apply_deliverable_coverage_audit,
    extract_deliverable_draft_text,
    render_deliverables,
)


class ArtifactRenderingTests(unittest.TestCase):
    def test_xlsx_renderer_uses_deliverable_contract_sheets(self) -> None:
        packet = {
            "deliverable_contract": {
                "deliverables": [
                    {
                        "filename": "section-382-analysis-workbook.xlsx",
                        "workbook_sheets": [
                            {"name": "Shareholder Register", "purpose": "Owner register"},
                            {"name": "Ownership Shift Calculations", "purpose": "Shift math"},
                            {"name": "Section 382 Limit Computation", "purpose": "Limit math"},
                        ],
                    }
                ]
            },
            "draft_answer": """
| Shareholder | Date | Percentage |
| --- | --- | --- |
| Alpha Fund | 2024-01-01 | 6.2% |
""",
            "verified_evidence": [
                {
                    "claim": "Alpha Fund exceeded five percent.",
                    "raw_support": "Alpha Fund held 6.2 percent.",
                    "source": {"doc_id": "doc_1", "chunk_id": "chunk_1"},
                }
            ],
        }
        with tempfile.TemporaryDirectory() as tmp:
            artifacts = render_deliverables(
                output_dir=tmp,
                deliverables=["section-382-analysis-workbook.xlsx"],
                title="Section 382",
                packet=packet,
            )
            path = Path(artifacts[0]["path"])
            workbook = load_workbook(path, read_only=True)
            try:
                self.assertIn("Shareholder Register", workbook.sheetnames)
                self.assertIn("Ownership Shift Calculations", workbook.sheetnames)
                self.assertIn("Section 382 Limit Computation", workbook.sheetnames)
                self.assertIn("Evidence", workbook.sheetnames)
                self.assertGreater(workbook["Shareholder Register"].max_row, 4)
            finally:
                workbook.close()

    def test_docx_renderer_replaces_encoded_artifact_with_readable_fallback(self) -> None:
        packet = {
            "draft_answer": "```xml\n<base64_file><filename>memo.docx</filename><content>"
            + ("A" * 5000)
            + "</content></base64_file>\n```",
            "cheap_worker_summary": "Class 1 interest underpayment and Class 6 cash pool issue.",
            "verified_evidence": [],
        }
        with tempfile.TemporaryDirectory() as tmp:
            artifacts = render_deliverables(
                output_dir=tmp,
                deliverables=["memo.docx"],
                title="Memo",
                packet=packet,
            )
            doc = Document(Path(artifacts[0]["path"]))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            self.assertIn("Class 1 interest underpayment", text)
            self.assertNotIn("<base64_file>", text)

    def test_docx_renderer_preserves_structured_appendix(self) -> None:
        packet = {
            "draft_answer": "Clean memo narrative.",
            "artifact_appendix": "Score-critical row: 2023 Merger Guidelines threshold.",
            "verified_evidence": [],
        }
        with tempfile.TemporaryDirectory() as tmp:
            artifacts = render_deliverables(
                output_dir=tmp,
                deliverables=["memo.docx"],
                title="Memo",
                packet=packet,
            )
            doc = Document(Path(artifacts[0]["path"]))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            self.assertIn("Structured Findings Appendix", text)
        self.assertIn("2023 Merger Guidelines threshold", text)

    def test_docx_renderer_adds_missing_deliverable_coverage_fill(self) -> None:
        packet = {
            "deliverable_contract": {
                "deliverables": [
                    {
                        "filename": "seller-certificate.docx",
                        "artifact_role": "seller_certificate",
                        "artifact_goal": "Draft the seller certificate.",
                        "required_sections": [
                            "Seller certificate title",
                            "Bringdown representations and warranties",
                        ],
                    },
                    {
                        "filename": "mac-certificate.docx",
                        "artifact_role": "mac_certificate",
                        "artifact_goal": "Draft the no-MAC certificate.",
                        "required_sections": ["Covered period and no-MAC statement"],
                    },
                ]
            },
            "draft_answer": "mac-certificate.docx\n\nCovered period and no-MAC statement.",
            "cheap_worker_summary": "Seller certificate should bring down the reps and disclose schedule exceptions.",
            "verified_evidence": [],
        }
        audit = apply_deliverable_coverage_audit(
            packet,
            ["seller-certificate.docx", "mac-certificate.docx"],
        )
        self.assertFalse(audit[0]["present_in_draft"])
        self.assertTrue(audit[1]["present_in_draft"])
        with tempfile.TemporaryDirectory() as tmp:
            artifacts = render_deliverables(
                output_dir=tmp,
                deliverables=["seller-certificate.docx", "mac-certificate.docx"],
                title="Disclosure Package",
                packet=packet,
            )
            seller_doc = Document(Path(artifacts[0]["path"]))
            seller_text = "\n".join(paragraph.text for paragraph in seller_doc.paragraphs)
            mac_doc = Document(Path(artifacts[1]["path"]))
            mac_text = "\n".join(paragraph.text for paragraph in mac_doc.paragraphs)
        self.assertIn("Deliverable-Specific Coverage Fill", seller_text)
        self.assertIn("Bringdown representations and warranties", seller_text)
        self.assertIn("Seller certificate should bring down the reps", seller_text)
        self.assertNotIn("Deliverable-Specific Coverage Fill", mac_text)

    def test_docx_renderer_writes_matching_filename_section_to_each_deliverable(self) -> None:
        packet = {
            "draft_answer": """# seller-certificate.docx

Seller-only certificate text.

# mac-certificate.docx

MAC-only certificate text.
""",
            "verified_evidence": [],
        }
        with tempfile.TemporaryDirectory() as tmp:
            artifacts = render_deliverables(
                output_dir=tmp,
                deliverables=["seller-certificate.docx", "mac-certificate.docx"],
                title="Disclosure Package",
                packet=packet,
            )
            seller_doc = Document(Path(artifacts[0]["path"]))
            seller_text = "\n".join(paragraph.text for paragraph in seller_doc.paragraphs)
            mac_doc = Document(Path(artifacts[1]["path"]))
            mac_text = "\n".join(paragraph.text for paragraph in mac_doc.paragraphs)
        self.assertIn("Seller-only certificate text", seller_text)
        self.assertNotIn("MAC-only certificate text", seller_text)
        self.assertIn("MAC-only certificate text", mac_text)
        self.assertNotIn("Seller-only certificate text", mac_text)

    def test_deliverable_section_extraction_preserves_global_draft_without_heading(self) -> None:
        text = "One global answer without exact filename sections."
        self.assertEqual(
            extract_deliverable_draft_text(
                text,
                deliverable="memo.docx",
                deliverables=["memo.docx", "schedule.docx"],
            ),
            text,
        )

    def test_xlsx_renderer_keeps_cells_evaluator_sized(self) -> None:
        packet = {
            "draft_answer": "| Issue | Finding |\n| --- | --- |\n| A | " + ("long " * 300) + " |",
            "verified_evidence": [
                {
                    "claim": "A" * 800,
                    "raw_support": "B" * 1200,
                    "source": {"doc_id": "doc_1", "chunk_id": "chunk_1"},
                }
            ],
        }
        with tempfile.TemporaryDirectory() as tmp:
            artifacts = render_deliverables(
                output_dir=tmp,
                deliverables=["matrix.xlsx"],
                title="Matrix",
                packet=packet,
            )
            workbook = load_workbook(Path(artifacts[0]["path"]), read_only=True)
            try:
                for sheet in workbook.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        for value in row:
                            if isinstance(value, str):
                                self.assertLessEqual(len(value), 500)
            finally:
                workbook.close()

    def test_xlsx_renderer_prefers_local_plain_heading_table(self) -> None:
        packet = {
            "deliverables": ["clawback-candidate-list.xlsx", "memo.docx"],
            "deliverable_contract": {
                "deliverables": [
                    {
                        "filename": "clawback-candidate-list.xlsx",
                        "artifact_role": "candidate_list_workbook",
                        "workbook_sheets": [
                            {
                                "name": "Clawback Candidates",
                                "purpose": "Privilege-log entry, Bates range, claimed basis, deficiency, recommended action, and source support",
                            }
                        ],
                    }
                ]
            },
            "draft_answer": """
# clawback-candidate-list.xlsx

Workbook overview.

Clawback Candidates
| Privilege Log Entry / Doc ID | Bates Range | Privilege Basis Claimed | Deficiency / Gap | Recommended Action | Source ID |
| --- | --- | --- | --- | --- | --- |
| 24 | TF-PRIV-000190-197 | ACP | Non-lawyer-only operational communication | Produce or reclassify | doc_0024 |

# memo.docx

Narrative memo.
""",
            "cheap_worker_summary": """
### **1. Fact Allocation Inventory**
| Source Document | Fact / Issue | Deliverable Role |
| --- | --- | --- |
| doc_0001 | Generic background | Memo |
""",
            "verified_evidence": [],
        }
        with tempfile.TemporaryDirectory() as tmp:
            artifacts = render_deliverables(
                output_dir=tmp,
                deliverables=["clawback-candidate-list.xlsx"],
                title="Privilege Review",
                packet=packet,
            )
            workbook = load_workbook(Path(artifacts[0]["path"]), read_only=True)
            try:
                rows = list(workbook["Clawback Candidates"].iter_rows(values_only=True))
            finally:
                workbook.close()
        flattened = "\n".join(" | ".join(str(cell or "") for cell in row) for row in rows)
        self.assertIn("Privilege Log Entry / Doc ID", flattened)
        self.assertIn("Non-lawyer-only operational communication", flattened)
        self.assertNotIn("Generic background", flattened)


if __name__ == "__main__":
    unittest.main()
